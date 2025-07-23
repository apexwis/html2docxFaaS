import os
from flask import Flask, request, send_file, jsonify
from docx import Document
from bs4 import BeautifulSoup
import tempfile
import traceback
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

app = Flask(__name__)

API_KEY = os.environ.get('API_KEY')

def require_api_key():
    auth = request.headers.get('Authorization', '')
    if not auth.startswith('Bearer '):
        return False
    token = auth.split(' ', 1)[1]
    return token == API_KEY

# Example: Standardize the DOCX structure
# You can expand this function to match your protocol's needs
def html_to_standardized_docx(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    doc = Document()

    # Add logo to header
    logo_path = 'logo_kontiki.png'
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align to right
    run = paragraph.add_run()
    run.add_picture(logo_path, width=Inches(0.7))  # Adjust width as needed

    body = soup.body

    for elem in body.children:
        if elem.name is None:
            continue  # Skip text nodes or whitespace
        if elem.name in ['h1', 'h2', 'h3']:
            level = {'h1': 0, 'h2': 1, 'h3': 2}[elem.name]
            heading = doc.add_heading(elem.get_text(), level=level)
            # Set heading font to Arial
            for run in heading.runs:
                run.font.name = 'Arial'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        elif elem.name == 'p':
            para = doc.add_paragraph(elem.get_text())
            # Set body text font to Arial 11
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        elif elem.name == 'table':
            rows = elem.find_all('tr')
            if not rows:
                continue
            cols = rows[0].find_all(['td', 'th'])
            table_docx = doc.add_table(rows=len(rows), cols=len(cols))
            table_docx.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    cell_obj = table_docx.cell(i, j)
                    cell_obj.text = cell.get_text()
                    # Set table cell font to Arial 11
                    for paragraph in cell_obj.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(11)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return _save_docx_to_tempfile(doc)

def _save_docx_to_tempfile(doc):
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()
    print('DOCX created at', tmp.name)
    return tmp.name

@app.before_request
def log_request_info():
    print(f"Received {request.method} request for {request.url}")
    print(f"Headers: {dict(request.headers)}")
    print(f"Body: {request.get_data(as_text=True)[:1000]}")  # Print up to 1000 chars

@app.route('/convert', methods=['POST'])
def convert():
    try:
        if not require_api_key():
            print('Unauthorized request')
            return jsonify({'error': 'Unauthorized'}), 401
        if not request.data:
            print('No HTML provided in request')
            return jsonify({'error': 'No HTML provided'}), 400
        html_content = request.data.decode('utf-8')
        docx_path = html_to_standardized_docx(html_content)
        response = send_file(docx_path, as_attachment=True, download_name='protocol.docx')
        @response.call_on_close
        def cleanup():
            os.remove(docx_path)
        return response
    except Exception as e:
        print('Exception occurred:')
        print(traceback.format_exc())
        return jsonify({'error': 'Internal server error', 'details': str(e), 'trace': traceback.format_exc()}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0')
